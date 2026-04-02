#!/usr/bin/env python3
"""
Batch E2E Test Runner for GYN Tumor Board

Processes all 15 real patient GUIDs through the full 10-agent group chat workflow.
Validates artifacts (Word doc + PPTX) and produces a pass/fail summary report.

Usage:
    cd src
    python3 ../scripts/run_batch_e2e.py [--patients N] [--timeout SECS] [--print]

Prerequisites:
    - Azure OpenAI credentials in .env (API key or `az login`)
    - Patient data CSVs in ../infra/patient_data/
    - CLINICAL_NOTES_SOURCE=caboodle in .env
"""

import argparse
import asyncio
import csv
import json
import logging
import os
import re
import sys
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

# Ensure src/ is on the path
SRC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "src")
sys.path.insert(0, SRC_DIR)

from dotenv import load_dotenv

load_dotenv(os.path.join(SRC_DIR, ".env"))

# Force caboodle mode for local CSV data
os.environ.setdefault("CLINICAL_NOTES_SOURCE", "caboodle")
os.environ.setdefault("SCENARIO", "default")
os.environ.setdefault("EXCLUDED_AGENTS", "")

from config import load_agent_config
from data_models.app_context import AppContext
from data_models.chat_context import ChatContext
from data_models.epic.caboodle_file_accessor import CaboodleFileAccessor
from evaluation.chat_simulator import ChatSimulator, ProceedUser
from group_chat import create_group_chat
from tests.local_accessors import create_local_data_access

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

DATA_DIR = os.path.join(SRC_DIR, "..", "infra", "patient_data")
QUERIES_CSV = os.path.join(SRC_DIR, "evaluation", "initial_queries_gyn15.csv")
OUTPUT_DIR = os.path.join(SRC_DIR, "evaluation", "batch_e2e_output")
DEFAULT_TIMEOUT = 300  # 5 minutes per patient
MAX_TURNS = 5  # Max user turns (each turn may have many agent messages)

# ICD-10 codes for NCCN-covered GYN cancers
NCCN_COVERED_ICD10_PREFIXES = {
    "C54": "uterine",     # Uterine corpus
    "C55": "uterine",     # Uterine NOS
    "C52": "vaginal",     # Vagina
    "C51": "vulvar",      # Vulva
}

# NCCN page code prefixes to look for in agent output
NCCN_PAGE_PREFIXES = ["ENDO-", "UTSARC-", "VAG-", "VULVA-", "VM-"]


@dataclass
class PatientResult:
    """Result of processing a single patient through the full workflow."""
    patient_id: str
    status: str  # PASS, FAIL, TIMEOUT, SKIP
    duration_secs: float
    agents_responded: list[str]
    agents_total: int
    docx_ok: bool
    pptx_ok: bool
    nccn_codes_cited: list[str]
    nccn_applicable: bool
    disease_site: str
    error: str | None


# ---------------------------------------------------------------------------
# AppContext factory for local testing
# ---------------------------------------------------------------------------

def create_local_app_context() -> AppContext:
    """Create an AppContext suitable for local batch testing.

    Uses API key auth if AZURE_OPENAI_API_KEY is set, otherwise AzureCliCredential.
    Uses local file accessors instead of Azure Blob Storage.
    """
    agent_configs = load_agent_config(os.environ.get("SCENARIO", "default"))
    data_access = create_local_data_access(data_dir=DATA_DIR, output_dir=OUTPUT_DIR)

    api_key = os.environ.get("AZURE_OPENAI_API_KEY")
    if api_key:
        # API key auth — no credential needed for OpenAI calls
        # AppContext requires credential, pass None and rely on API key path in group_chat.py
        return AppContext(
            all_agent_configs=agent_configs,
            blob_service_client=None,
            credential=None,
            data_access=data_access,
        )
    else:
        # Token-based auth via Azure CLI
        from azure.identity.aio import AzureCliCredential
        credential = AzureCliCredential()
        return AppContext(
            all_agent_configs=agent_configs,
            blob_service_client=None,
            credential=credential,
            data_access=data_access,
        )


# ---------------------------------------------------------------------------
# Artifact validation
# ---------------------------------------------------------------------------

def validate_docx(output_dir: str, conversation_id: str, patient_id: str) -> bool:
    """Check if a Word doc was generated and has content."""
    artifact_dir = os.path.join(output_dir, conversation_id, patient_id)
    if not os.path.isdir(artifact_dir):
        # Also check flat output dir
        artifact_dir = output_dir

    docx_files = list(Path(artifact_dir).rglob("*.docx"))
    if not docx_files:
        logger.warning("No .docx file found for %s", patient_id)
        return False

    try:
        from docx import Document
        doc = Document(str(docx_files[0]))
        # Check that the document has content (tables or paragraphs)
        has_content = len(doc.tables) > 0 or any(p.text.strip() for p in doc.paragraphs)
        if has_content:
            logger.info("DOCX OK: %s (%d tables, %d paragraphs)", docx_files[0].name, len(doc.tables), len(doc.paragraphs))
        return has_content
    except Exception as e:
        logger.warning("Failed to validate .docx for %s: %s", patient_id, e)
        return False


def validate_pptx(output_dir: str, conversation_id: str, patient_id: str) -> bool:
    """Check if a PPTX was generated with 3 slides."""
    artifact_dir = os.path.join(output_dir, conversation_id, patient_id)
    if not os.path.isdir(artifact_dir):
        artifact_dir = output_dir

    pptx_files = list(Path(artifact_dir).rglob("*.pptx"))
    if not pptx_files:
        logger.warning("No .pptx file found for %s", patient_id)
        return False

    try:
        from pptx import Presentation
        prs = Presentation(str(pptx_files[0]))
        slide_count = len(prs.slides)
        logger.info("PPTX OK: %s (%d slides)", pptx_files[0].name, slide_count)
        return slide_count >= 3
    except Exception as e:
        logger.warning("Failed to validate .pptx for %s: %s", patient_id, e)
        return False


# ---------------------------------------------------------------------------
# Disease site detection and NCCN citation check
# ---------------------------------------------------------------------------

async def get_disease_site(patient_id: str) -> tuple[str, bool]:
    """Determine the patient's cancer type and whether NCCN tool applies.

    Returns (disease_site, nccn_applicable).
    """
    caboodle = CaboodleFileAccessor(data_dir=DATA_DIR)
    try:
        diagnoses = await caboodle.get_diagnoses(patient_id)
    except Exception:
        return "unknown", False

    for dx in diagnoses:
        icd = dx.get("ICD10Code", "")
        for prefix, site in NCCN_COVERED_ICD10_PREFIXES.items():
            if icd.startswith(prefix):
                return site, True

    # Check for ovarian (C56), cervical (C53), etc.
    for dx in diagnoses:
        icd = dx.get("ICD10Code", "")
        if icd.startswith("C56"):
            return "ovarian", False
        if icd.startswith("C53"):
            return "cervical", False

    return "other_gyn", False


def extract_nccn_citations(chat_text: str) -> list[str]:
    """Extract NCCN page code citations from agent chat output."""
    codes = set()
    for prefix in NCCN_PAGE_PREFIXES:
        pattern = rf"{prefix}\d+[A-Z]?"
        matches = re.findall(pattern, chat_text.upper())
        codes.update(matches)
    return sorted(codes)


def extract_agents_from_chat(chat_text: str, all_agent_names: list[str]) -> list[str]:
    """Identify which agents produced output in the chat history."""
    responded = []
    for name in all_agent_names:
        # Readable format: "ASSISTANT (agent id: Orchestrator):"
        # Or direct: "Orchestrator:" at start of line
        if f"agent id: {name}" in chat_text or re.search(rf"^{re.escape(name)}:", chat_text, re.MULTILINE):
            responded.append(name)
    return responded


# ---------------------------------------------------------------------------
# Single patient runner
# ---------------------------------------------------------------------------

async def run_single_patient(
    app_ctx: AppContext,
    patient_id: str,
    initial_query: str,
    timeout: int,
    print_messages: bool = False,
) -> PatientResult:
    """Run a single patient through the full group chat workflow."""

    disease_site, nccn_applicable = await get_disease_site(patient_id)
    agent_names = [a["name"] for a in app_ctx.all_agent_configs]

    logger.info("=" * 60)
    logger.info("Patient: %s | Disease: %s | NCCN: %s", patient_id[:8] + "...", disease_site, nccn_applicable)
    logger.info("=" * 60)

    start = time.time()
    chat_text = ""
    status = "FAIL"
    error = None

    try:
        # Create group chat
        chat_ctx = ChatContext(f"batch-e2e-{patient_id[:8]}")
        chat_ctx.patient_id = patient_id
        group_chat, chat_ctx = create_group_chat(app_ctx, chat_ctx)

        # Create simulator
        simulated_user = ProceedUser()
        simulated_user.setup(patient_id, initial_query)

        simulator = ChatSimulator(
            simulated_user=simulated_user,
            group_chat_kwargs={"app_ctx": app_ctx},
            patients_id=[patient_id],
            initial_queries=[initial_query],
            trial_count=1,
            max_turns=MAX_TURNS,
            output_folder_path=OUTPUT_DIR,
            save_readable_history=True,
            print_messages=print_messages,
        )
        # Reuse the already-created group chat
        simulator.group_chat = group_chat
        simulator.chat_context = chat_ctx

        # Run with timeout
        await asyncio.wait_for(
            simulator.chat(patient_id, initial_query, [], MAX_TURNS),
            timeout=timeout,
        )

        # Save chat history
        safe_id = patient_id[:8]
        simulator.save(
            f"chat_{safe_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            save_readable_history=True,
        )

        # Extract readable chat text for analysis
        from evaluation.utils import chat_history_to_readable_text
        chat_text = chat_history_to_readable_text(group_chat.history)

        status = "PASS"

    except asyncio.TimeoutError:
        status = "TIMEOUT"
        error = f"Exceeded {timeout}s timeout"
        logger.error("TIMEOUT for patient %s after %ds", patient_id[:8], timeout)
    except Exception as e:
        status = "FAIL"
        error = str(e)[:200]
        logger.error("FAIL for patient %s: %s", patient_id[:8], e)

    duration = time.time() - start

    # Post-run validation
    conversation_id = f"batch-e2e-{patient_id[:8]}"
    docx_ok = validate_docx(OUTPUT_DIR, conversation_id, patient_id) if status == "PASS" else False
    pptx_ok = validate_pptx(OUTPUT_DIR, conversation_id, patient_id) if status == "PASS" else False
    agents_responded = extract_agents_from_chat(chat_text, agent_names) if chat_text else []
    nccn_codes = extract_nccn_citations(chat_text) if chat_text and nccn_applicable else []

    return PatientResult(
        patient_id=patient_id,
        status=status,
        duration_secs=duration,
        agents_responded=agents_responded,
        agents_total=len(agent_names),
        docx_ok=docx_ok,
        pptx_ok=pptx_ok,
        nccn_codes_cited=nccn_codes,
        nccn_applicable=nccn_applicable,
        disease_site=disease_site,
        error=error,
    )


# ---------------------------------------------------------------------------
# Summary report
# ---------------------------------------------------------------------------

def print_summary(results: list[PatientResult]):
    """Print a formatted summary table to console."""
    print()
    print("=" * 100)
    print(f"BATCH E2E TEST REPORT — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 100)
    print()
    print(f"{'Patient GUID':<42} {'Status':<9} {'Time':<8} {'Agents':<9} {'Doc':<5} {'PPTX':<5} {'NCCN'}")
    print("-" * 100)

    pass_count = 0
    fail_count = 0
    timeout_count = 0

    for r in results:
        guid_short = r.patient_id
        time_str = f"{r.duration_secs:.0f}s"
        agents_str = f"{len(r.agents_responded)}/{r.agents_total}"
        doc_str = "OK" if r.docx_ok else "MISS"
        pptx_str = "OK" if r.pptx_ok else "MISS"

        if r.nccn_applicable:
            nccn_str = ",".join(r.nccn_codes_cited[:3]) if r.nccn_codes_cited else "none"
        else:
            nccn_str = f"N/A ({r.disease_site})"

        status_str = r.status
        if r.status == "PASS":
            pass_count += 1
        elif r.status == "TIMEOUT":
            timeout_count += 1
        else:
            fail_count += 1

        print(f"{guid_short:<42} {status_str:<9} {time_str:<8} {agents_str:<9} {doc_str:<5} {pptx_str:<5} {nccn_str}")

        if r.error:
            print(f"  ERROR: {r.error}")

    print("-" * 100)
    total = len(results)
    print(f"TOTAL: {pass_count}/{total} PASS, {fail_count} FAIL, {timeout_count} TIMEOUT")
    print()


def save_summary_json(results: list[PatientResult], output_dir: str):
    """Save detailed results to JSON."""
    summary = {
        "timestamp": datetime.now().isoformat(),
        "total_patients": len(results),
        "passed": sum(1 for r in results if r.status == "PASS"),
        "failed": sum(1 for r in results if r.status == "FAIL"),
        "timed_out": sum(1 for r in results if r.status == "TIMEOUT"),
        "results": [
            {
                "patient_id": r.patient_id,
                "status": r.status,
                "duration_secs": round(r.duration_secs, 1),
                "agents_responded": r.agents_responded,
                "agents_total": r.agents_total,
                "docx_ok": r.docx_ok,
                "pptx_ok": r.pptx_ok,
                "nccn_codes_cited": r.nccn_codes_cited,
                "nccn_applicable": r.nccn_applicable,
                "disease_site": r.disease_site,
                "error": r.error,
            }
            for r in results
        ],
    }

    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"batch_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(path, "w") as f:
        json.dump(summary, f, indent=2)
    print(f"Summary saved to: {path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

async def main():
    parser = argparse.ArgumentParser(description="Batch E2E test runner for GYN Tumor Board")
    parser.add_argument("--patients", type=int, default=None, help="Number of patients to process (default: all 15)")
    parser.add_argument("--timeout", type=int, default=DEFAULT_TIMEOUT, help=f"Per-patient timeout in seconds (default: {DEFAULT_TIMEOUT})")
    parser.add_argument("--print", action="store_true", dest="print_messages", help="Print agent messages to console")
    parser.add_argument("--csv", type=str, default=QUERIES_CSV, help="Path to initial queries CSV")
    parser.add_argument("--patient-id", type=str, default=None, help="Run a single patient by GUID (overrides --csv)")
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
        datefmt="%H:%M:%S",
    )

    # Suppress noisy loggers
    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("semantic_kernel").setLevel(logging.WARNING)

    # Ensure BOT_IDS is set (required by config loader)
    os.environ.setdefault(
        "BOT_IDS",
        json.dumps({name: "dummy" for name in [
            "Orchestrator", "PatientHistory", "OncologicHistory", "Pathology",
            "Radiology", "PatientStatus", "ClinicalGuidelines", "ReportCreation",
            "ClinicalTrials", "MedicalResearch",
        ]})
    )
    os.environ.setdefault("HLS_MODEL_ENDPOINTS", "{}")

    # Create app context
    app_ctx = create_local_app_context()
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Load patient list
    if args.patient_id:
        patients = [(args.patient_id, f"Prepare tumor board case for patient {args.patient_id}")]
    else:
        patients = []
        with open(args.csv, encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                pid = row["patient_id"].strip()
                query = row["initial_query"].strip()
                if pid:
                    patients.append((pid, query))

    if args.patients:
        patients = patients[:args.patients]

    logger.info("Batch E2E: %d patients, timeout=%ds", len(patients), args.timeout)

    # Process each patient sequentially
    results: list[PatientResult] = []
    for i, (patient_id, initial_query) in enumerate(patients):
        logger.info("[%d/%d] Processing %s...", i + 1, len(patients), patient_id[:8])
        result = await run_single_patient(
            app_ctx=app_ctx,
            patient_id=patient_id,
            initial_query=initial_query,
            timeout=args.timeout,
            print_messages=args.print_messages,
        )
        results.append(result)
        logger.info("[%d/%d] %s — %s (%.0fs)", i + 1, len(patients), patient_id[:8], result.status, result.duration_secs)

    # Report
    print_summary(results)
    save_summary_json(results, OUTPUT_DIR)

    # Clean up credential if needed
    if app_ctx.credential is not None:
        await app_ctx.credential.close()

    # Exit code: 0 if all pass, 1 if any fail
    if all(r.status == "PASS" for r in results):
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
