#!/usr/bin/env python3
"""Generate the GYN Tumor Board landscape Word template programmatically.

Produces tumor_board_template.docx with:
  - Landscape orientation (11" x 8.5")
  - 5-column table matching the real Rush tumor board handout format:
    Col 0: Patient (case #, MRN, attending, RTC, location, path date)
    Col 1: Diagnosis & Pertinent History
    Col 2: Previous Tx or Operative Findings, Tumor Markers
    Col 3: Imaging
    Col 4: Discussion
  - Column widths: 0.93 / 2.00 / 2.94 / 3.31 / 1.12 inches
  - Margins: 0.5L / 0.5R / 0.6T / 0.5B (matches real doc exactly)
  - Jinja2 placeholders for docxtpl rendering

Run: python scripts/generate_docx_template.py
Output: src/scenarios/default/templates/tumor_board_template.docx
"""

import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


# Column definitions: (header text, width in inches)
# Widths measured from real Rush tumor board handout (TB Handout 03.04.2026.docx)
COLUMNS = [
    ("Patient", 0.93),
    ("Diagnosis & Pertinent History", 2.00),
    ("Previous Tx or Operative Findings, Tumor Markers", 2.94),
    ("Imaging", 3.31),
    ("Discussion", 1.12),
]

# Jinja2 placeholders per column (for docxtpl)
# Using {{r var}} for RichText support
CELL_PLACEHOLDERS = [
    "{{r col0_content}}",
    "{{r col1_content}}",
    "{{r col2_content}}",
    "{{r col3_content}}",
    "{{r col4_content}}",
]


def _set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_width(cell, width):
    """Set cell width."""
    cell.width = width


def _set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    """Set cell margins in twips (1/1440 inch)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def create_template():
    doc = Document()

    # --- Page setup: landscape ---
    # Margins match real Rush tumor board handout exactly
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # --- Create 5-column table (header row + content row) ---
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # --- Header row ---
    for i, (header_text, col_width) in enumerate(COLUMNS):
        cell = table.rows[0].cells[i]
        _set_cell_width(cell, Inches(col_width))
        _set_cell_margins(cell)

        # Clear default paragraph
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # --- Content row with jinja2 placeholders ---
    for i, (_, col_width) in enumerate(COLUMNS):
        cell = table.rows[1].cells[i]
        _set_cell_width(cell, Inches(col_width))
        _set_cell_margins(cell)

        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(CELL_PLACEHOLDERS[i])
        run.font.size = Pt(9)
        run.font.name = "Calibri"

    return doc


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    repo_root = os.path.dirname(script_dir)
    output_path = os.path.join(
        repo_root, "src", "scenarios", "default", "templates", "tumor_board_template.docx"
    )

    doc = create_template()
    doc.save(output_path)
    print(f"Template saved to: {output_path}")


if __name__ == "__main__":
    main()
